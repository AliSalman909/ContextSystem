import os
from pathlib import Path
from fastapi import UploadFile
from sqlalchemy.orm import Session

from docx import Document as DocxDocument
from pypdf import PdfReader

from app import db

from ..settings import settings
from .. import models
from .chunking import chunk_text

from .embeddings import embed_text
from .pinecone_client import index as pinecone_index

def _ensure_dir(path: str) -> None:
    Path(path).mkdir(parents=True, exist_ok=True)


def _save_upload(customer_id: str, upload_file: UploadFile) -> str:
    base_dir = settings.UPLOAD_DIR
    customer_dir = os.path.join(base_dir, customer_id)
    _ensure_dir(customer_dir)

    file_path = os.path.join(customer_dir, upload_file.filename)
    with open(file_path, "wb") as f:
        f.write(upload_file.file.read())
    return file_path


def _extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()

    if ext == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return text.strip()

    if ext == ".pdf":
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()

    try:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()
    except Exception:
        return ""


def ingest_document(db: Session, customer_id: str, doc_type: str, upload_file: UploadFile) -> models.Document:
    
    # 1) Save file
    file_path = _save_upload(customer_id, upload_file)

    # 2) Document row
    doc = models.Document(
        customer_id=customer_id,
        doc_type=doc_type,
        filename=upload_file.filename,
        storage_path=file_path,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # 3) Extract text
    extracted = _extract_text(file_path)
    if not extracted:
        extracted = "(No text extracted from this file. Try a .txt/.docx/.pdf with selectable text.)"

    # 4) Store extracted text
    doc_text = models.DocumentText(document_id=doc.id, extracted_text=extracted)
    db.add(doc_text)
    db.commit()

    # 5) Chunk and store chunks
    chunks = chunk_text(extracted, chunk_size=900, overlap=150)
    
    for i, ch in enumerate(chunks):
        vector = embed_text(ch)

        vector_id = f"{doc.id}_{i}"

        pinecone_index.upsert(
            vectors=[
                {
                    "id": vector_id,
                    "values": vector,
                    "metadata": {
                        "customer_id": customer_id,
                        "document_id": doc.id,
                        "chunk_index": i,
                        "doc_type": doc_type,
                        "uploaded_at": doc.uploaded_at.isoformat(), #newl1
                        "text": ch,                              #newl2
                    },
                }
            ]
        )

        row = models.Chunk(
            document_id=doc.id,
            chunk_index=i,
            chunk_text=ch,
            pinecone_vector_id=vector_id,
        )
        db.add(row)

    db.commit()
    

    return doc
